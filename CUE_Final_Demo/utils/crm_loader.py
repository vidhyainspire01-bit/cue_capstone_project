# utils/crm_loader.py
import json, uuid
from pathlib import Path
from docx import Document

def read_docx_rows(path: Path):
    doc = Document(path)
    # try to parse table rows first (common for exports)
    rows = []
    if doc.tables:
        for table in doc.tables:
            hdr = [c.text.strip() for c in table.rows[0].cells]
            for r in table.rows[1:]:
                row = {hdr[i]: r.cells[i].text.strip() for i in range(len(hdr))}
                rows.append(row)
    else:
        # fallback: paragraphs -> treat each paragraph as a row if it looks like key:value or CSV-like
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            # try splitting on ' - ' or ':' or '\t' or '|'
            if ':' in t:
                k, v = t.split(':', 1)
                rows.append({k.strip(): v.strip()})
            else:
                rows.append({"text": t})
    return rows

def emit_chunks_from_crm_file(path: Path):
    chunks=[]
    suf = path.suffix.lower()
    if suf == '.docx':
        rows = read_docx_rows(path)
        i = 0
        for r in rows:
            # normalize fields if exists
            note = r.get('CRM_Note') or r.get('Note') or r.get('text') or ' '.join([str(v) for v in r.values()])
            customer = r.get('Customer') or r.get('customer') or None
            rec = {
                "chunk_id": str(uuid.uuid4()),
                "doc_id": r.get('Record ID') or f"CRM-{path.stem}-{i}",
                "doc_type": "crm",
                "customer_id": customer,
                "date": r.get('Date') or None,
                "chunk_index": 0,
                "tags": [r.get('Theme')] if r.get('Theme') else [],
                "text": note,
                "source_path": str(path)
            }
            chunks.append(rec)
            i += 1
    else:
        # keep backward CSV/JSON support optionally
        import pandas as pd
        df = pd.read_csv(path) if suf=='.csv' else pd.read_json(path, lines=True)
        for _, row in df.iterrows():
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "doc_id": row.get('Record ID') or f"CRM-{_}",
                "doc_type": "crm",
                "customer_id": row.get('Customer'),
                "date": row.get('Date'),
                "chunk_index": 0,
                "tags": [row.get('Theme')] if row.get('Theme') else [],
                "text": row.get('CRM_Note') or row.get('text') or '',
                "source_path": str(path)
            })
    return chunks
