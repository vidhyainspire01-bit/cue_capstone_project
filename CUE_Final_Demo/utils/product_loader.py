# utils/product_loader.py
import json, uuid
from pathlib import Path
from docx import Document

def read_docx_table_rows(path: Path):
    doc = Document(path)
    rows = []
    if doc.tables:
        for table in doc.tables:
            hdr = [c.text.strip() for c in table.rows[0].cells]
            for r in table.rows[1:]:
                rows.append({hdr[i]: r.cells[i].text.strip() for i in range(len(hdr))})
    else:
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                rows.append({"text": t})
    return rows

def emit_chunks_from_product_usage(path: Path):
    chunks=[]
    suf = path.suffix.lower()
    if suf == '.docx':
        rows = read_docx_table_rows(path)
        i=0
        for r in rows:
            summary = r.get('Notes') or r.get('text') or ''
            # build short summary if structured fields exist
            if 'Active Users' in r or 'Sessions' in r:
                summary = f"Active Users: {r.get('Active Users','')}, Sessions: {r.get('Sessions','')}. Notes: {summary}"
            rec = {
                "chunk_id": str(uuid.uuid4()),
                "doc_id": r.get('Record ID') or f"PU-{path.stem}-{i}",
                "doc_type": "product",
                "customer_id": r.get('Customer') or None,
                "date": r.get('Period') or None,
                "chunk_index": 0,
                "tags": [r.get('Theme')] if r.get('Theme') else [],
                "text": summary,
                "source_path": str(path)
            }
            chunks.append(rec)
            i += 1
    else:
        import pandas as pd
        df = pd.read_csv(path) if suf=='.csv' else pd.read_json(path, lines=True)
        for _, row in df.iterrows():
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "doc_id": row.get('Record ID') or f"PU-{_}",
                "doc_type": "product",
                "customer_id": row.get('Customer'),
                "date": row.get('Period'),
                "chunk_index": 0,
                "tags": [row.get('Theme')] if row.get('Theme') else [],
                "text": row.get('Notes') or '',
                "source_path": str(path)
            })
    return chunks
