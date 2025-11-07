# utils/slack_loader.py
import uuid
from pathlib import Path
from docx import Document
import csv
import io

def _parse_table_rows(table):
    hdr = [c.text.strip() for c in table.rows[0].cells]
    rows = []
    for r in table.rows[1:]:
        row = {}
        for i, cell in enumerate(r.cells):
            key = hdr[i] if i < len(hdr) else f"col_{i}"
            row[key] = cell.text.strip()
        rows.append(row)
    return rows

def _parse_paragraph_tsv(paragraphs):
    """
    Some docx exports put tabular data into paragraphs separated by tabs.
    Detect header line (tabs present) and parse following lines.
    """
    lines = [p.text.strip() for p in paragraphs if p.text and p.text.strip()]
    if not lines:
        return []
    # detect header with tabs or multiple columns separated by '\t' or multiple spaces + separator
    header_line = None
    for i, ln in enumerate(lines[:3]):  # check first few lines for header
        if '\t' in ln or '    ' in ln or ' | ' in ln:
            header_line = ln
            start_idx = i + 1
            break
    if not header_line:
        return []
    # detect delimiter
    if '\t' in header_line:
        delim = '\t'
    elif ' | ' in header_line:
        delim = ' | '
    else:
        delim = None

    headers = [h.strip() for h in header_line.split(delim)] if delim else [h.strip() for h in header_line.split()]
    rows = []
    for ln in lines[start_idx:]:
        if not ln:
            continue
        if delim:
            parts = [p.strip() for p in ln.split(delim)]
        else:
            parts = [p.strip() for p in ln.split(None, len(headers)-1)]
        if len(parts) < len(headers):
            # pad
            parts = parts + [''] * (len(headers)-len(parts))
        row = {headers[i]: parts[i] for i in range(len(headers))}
        rows.append(row)
    return rows

def emit_chunks_from_slack_file(path: Path):
    chunks = []
    suf = path.suffix.lower()
    if suf == ".docx":
        doc = Document(path)

        # 1) tables
        any_rows = False
        for table in doc.tables:
            try:
                rows = _parse_table_rows(table)
                any_rows = any_rows or len(rows) > 0
                for i, r in enumerate(rows):
                    text = r.get("Message Summary") or r.get("Message") or r.get("text") or " | ".join(r.values())
                    chunks.append({
                        "chunk_id": str(uuid.uuid4()),
                        "doc_id": r.get("Message ID") or f"SLK-{path.stem}-{i}",
                        "doc_type": "slack",
                        "customer_id": r.get("Customer"),
                        "date": r.get("Date"),
                        "chunk_index": i,
                        "tags": [r.get("Theme")] if r.get("Theme") else [],
                        "text": text,
                        "metadata": r,
                        "source_path": str(path)
                    })
            except Exception:
                continue

        # 2) If no tables found, attempt TSV/paragraph-style parsing
        if not any_rows:
            rows = _parse_paragraph_tsv(doc.paragraphs)
            if rows:
                for i, r in enumerate(rows):
                    text = r.get("Message Summary") or r.get("Message") or r.get("text") or " | ".join(r.values())
                    chunks.append({
                        "chunk_id": str(uuid.uuid4()),
                        "doc_id": r.get("Message ID") or f"SLK-{path.stem}-{i}",
                        "doc_type": "slack",
                        "customer_id": r.get("Customer"),
                        "date": r.get("Date"),
                        "chunk_index": i,
                        "tags": [r.get("Theme")] if r.get("Theme") else [],
                        "text": text,
                        "metadata": r,
                        "source_path": str(path)
                    })

        # 3) fallback: use paragraphs (one per chunk)
        if not chunks:
            for i, p in enumerate(doc.paragraphs):
                t = p.text.strip()
                if not t:
                    continue
                chunks.append({
                    "chunk_id": str(uuid.uuid4()),
                    "doc_id": f"SLK-{path.stem}-{i}",
                    "doc_type": "slack",
                    "customer_id": None,
                    "date": None,
                    "chunk_index": i,
                    "tags": [],
                    "text": t,
                    "metadata": {},
                    "source_path": str(path)
                })

    elif suf == ".csv":
        import pandas as pd
        df = pd.read_csv(path)
        for idx, row in df.iterrows():
            text = row.get("Message Summary") or row.get("text") or ""
            metadata = {k: v for k, v in row.to_dict().items()}
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "doc_id": row.get("Message ID") or f"SLK-{path.stem}-{idx}",
                "doc_type": "slack",
                "customer_id": row.get("Customer"),
                "date": row.get("Date"),
                "chunk_index": int(idx),
                "tags": [row.get("Theme")] if row.get("Theme") else [],
                "text": str(text),
                "metadata": metadata,
                "source_path": str(path)
            })
    elif suf == ".json":
        import json
        data = json.load(open(path, 'r', encoding='utf-8'))
        # either list of records or dict
        records = data if isinstance(data, list) else data.get("messages", [])
        for idx, r in enumerate(records):
            text = r.get("text") or r.get("Message Summary") or ""
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "doc_id": r.get("Message ID") or f"SLK-{path.stem}-{idx}",
                "doc_type": "slack",
                "customer_id": r.get("Customer"),
                "date": r.get("Date"),
                "chunk_index": idx,
                "tags": [r.get("Theme")] if r.get("Theme") else [],
                "text": text,
                "metadata": r,
                "source_path": str(path)
            })
    else:
        # unknown extension: attempt to open and treat as plain text with TSV rows
        txt = path.read_text(encoding='utf-8', errors='ignore')
        lines = [l.strip() for l in txt.splitlines() if l.strip()]
        if not lines:
            return []
        # detect header
        if '\t' in lines[0] or '|' in lines[0] or ',' in lines[0]:
            # attempt CSV parse
            dialect = csv.Sniffer().sniff(lines[0])
            reader = csv.DictReader(io.StringIO("\n".join(lines)), dialect=dialect)
            for idx, row in enumerate(reader):
                text = row.get("Message Summary") or row.get("text") or ""
                chunks.append({
                    "chunk_id": str(uuid.uuid4()),
                    "doc_id": row.get("Message ID") or f"SLK-{path.stem}-{idx}",
                    "doc_type": "slack",
                    "customer_id": row.get("Customer"),
                    "date": row.get("Date"),
                    "chunk_index": idx,
                    "tags": [row.get("Theme")] if row.get("Theme") else [],
                    "text": text,
                    "metadata": row,
                    "source_path": str(path)
                })
        else:
            for i, l in enumerate(lines):
                chunks.append({
                    "chunk_id": str(uuid.uuid4()),
                    "doc_id": f"SLK-{path.stem}-{i}",
                    "doc_type": "slack",
                    "customer_id": None,
                    "date": None,
                    "chunk_index": i,
                    "tags": [],
                    "text": l,
                    "metadata": {},
                    "source_path": str(path)
                })
    return chunks
