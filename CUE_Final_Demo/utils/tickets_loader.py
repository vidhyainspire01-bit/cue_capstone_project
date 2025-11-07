# utils/tickets_loader.py
import json, uuid
from pathlib import Path
from docx import Document

def read_docx_ticket_entries(path: Path):
    doc = Document(path)
    rows=[]
    if doc.tables:
        for t in doc.tables:
            hdr=[c.text.strip() for c in t.rows[0].cells]
            for r in t.rows[1:]:
                rows.append({hdr[i]: r.cells[i].text.strip() for i in range(len(hdr))})
    else:
        for p in doc.paragraphs:
            t=p.text.strip()
            if t:
                # treat each para as a ticket summary line
                rows.append({"Summary": t})
    return rows

def emit_chunks_from_ticket_file(path: Path):
    chunks=[]
    suf = path.suffix.lower()
    if suf == '.docx':
        rows = read_docx_ticket_entries(path)
        i=0
        for r in rows:
            summary = r.get('Summary') or r.get('text') or ''
            rec = {
                "chunk_id": str(uuid.uuid4()),
                "doc_id": r.get('Ticket ID') or f"T-{path.stem}-{i}",
                "doc_type": "ticket",
                "customer_id": r.get('Customer') or None,
                "date": r.get('Date') or None,
                "chunk_index": 0,
                "tags": [r.get('Theme')] if r.get('Theme') else [],
                "text": summary,
                "source_path": str(path)
            }
            chunks.append(rec)
            i+=1
    else:
        import pandas as pd
        df = pd.read_csv(path) if suf=='.csv' else pd.read_json(path, lines=True)
        for _, row in df.iterrows():
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "doc_id": row.get('Ticket ID') or f"T-{_}",
                "doc_type": "ticket",
                "customer_id": row.get('Customer'),
                "date": row.get('Date'),
                "chunk_index": 0,
                "tags": [row.get('Theme')] if row.get('Theme') else [],
                "text": row.get('Summary') or row.get('body') or '',
                "source_path": str(path)
            })
    return chunks
